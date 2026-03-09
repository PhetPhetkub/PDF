from fastapi import FastAPI, File, UploadFile, Form
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pdf2docx import Converter
import pdfplumber
from deep_translator import GoogleTranslator
from docx import Document
import os
import uuid

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], 
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.post("/upload")
async def process_pdf(file: UploadFile = File(...), action: str = Form(...)):
    pdf_path = f"{uuid.uuid4()}.pdf"
    docx_path = f"{uuid.uuid4()}.docx"
    
    with open(pdf_path, "wb") as f:
        f.write(await file.read())

    try:
        if action == "convert":
            # 🟢 โหมดที่ 1: แปลงปกติ
            cv = Converter(pdf_path)
            cv.convert(docx_path)
            cv.close()
            output_filename = "Converted_Original.docx"

        elif action == "translate":
            # 🔵 โหมดที่ 2: แปลภาษา (อัปเกรดความเร็ว แปลรวดเดียวทั้งหน้า)
            doc = Document()
            doc.add_heading('เอกสารแปลภาษาไทย', 0)
            translator = GoogleTranslator(source='auto', target='th')
            
            with pdfplumber.open(pdf_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    doc.add_heading(f'--- หน้าที่ {i+1} ---', level=1)
                    text = page.extract_text()
                    
                    # ถอดรหัสแปลทีเดียวทั้งหน้า (ประหยัดเวลาเซิร์ฟเวอร์)
                    if text and text.strip():
                        try:
                            translated_text = translator.translate(text)
                            doc.add_paragraph(translated_text)
                        except Exception as e:
                            doc.add_paragraph(f"[แปลขัดข้อง: ข้อความอาจจะยาวเกินไป หรือเซิร์ฟเวอร์ Google ไม่ตอบสนอง]")
                    else:
                        doc.add_paragraph("[ไม่มีข้อความในหน้านี้ หรือเป็นรูปภาพล้วน]")
                        
                    doc.add_page_break()
            
            doc.save(docx_path)
            output_filename = "Translated_TH.docx"

        return FileResponse(
            docx_path, 
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
            filename=output_filename
        )
        
    finally:
        if os.path.exists(pdf_path):
            os.remove(pdf_path)
